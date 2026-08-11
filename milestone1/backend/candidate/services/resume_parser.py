import os
import re

import pymupdf
import spacy
from docx import Document


# ============================================================
# LOAD NLP MODEL
# ============================================================

nlp = spacy.load("en_core_web_sm")


# ============================================================
# COMMON TECHNICAL SKILLS
# ============================================================

SKILLS = [
    "python",
    "java",
    "c",
    "c++",
    "c#",
    "javascript",
    "typescript",
    "sql",
    "html",
    "css",
    "react",
    "angular",
    "vue",
    "node.js",
    "express",
    "django",
    "flask",
    "fastapi",
    "spring boot",
    "aws",
    "azure",
    "gcp",
    "docker",
    "kubernetes",
    "git",
    "github",
    "mongodb",
    "mysql",
    "postgresql",
    "oracle",
    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
    "machine learning",
    "deep learning",
    "data structures",
    "algorithms",
    "oop",
    "object oriented programming",
    "jdbc",
    "servlets",
]


# ============================================================
# TEXT EXTRACTION
# ============================================================

def extract_text_from_pdf(file_path):
    """Extract text from a PDF resume."""

    document = pymupdf.open(file_path)

    text = []

    for page in document:
        page_text = page.get_text()

        if page_text.strip():
            text.append(page_text)

    document.close()

    return "\n".join(text)


def extract_text_from_docx(file_path):
    """Extract text from a DOCX resume."""

    document = Document(file_path)

    text = []

    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text.append(paragraph.text)

    # Also extract text from tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():
                    text.append(cell.text)

    return "\n".join(text)


def extract_resume_text(file_path):
    """Choose the correct parser based on file extension."""

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)

    if extension == ".docx":
        return extract_text_from_docx(file_path)

    raise ValueError(
        "Unsupported file format. Only PDF and DOCX are supported."
    )


# ============================================================
# EMAIL
# ============================================================

def extract_email(text):
    """Extract email address from resume."""

    pattern = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"

    match = re.search(pattern, text)

    if match:
        return match.group(0)

    return None


# ============================================================
# PHONE
# ============================================================

def extract_phone(text):
    """Extract Indian phone number from resume."""

    patterns = [
        r"\+91[\s-]?[6-9]\d{9}",
        r"\b[6-9]\d{9}\b",
    ]

    for pattern in patterns:

        match = re.search(pattern, text)

        if match:
            return match.group(0)

    return None


# ============================================================
# NAME
# ============================================================

def extract_name(text):
    """Try to identify the candidate name."""

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    if not lines:
        return None

    ignored_words = [
        "resume",
        "curriculum vitae",
        "cv",
        "profile",
        "summary",
    ]

    for line in lines[:8]:

        clean_line = re.sub(
            r"[^\w\s.-]",
            "",
            line
        ).strip()

        if not clean_line:
            continue

        if len(clean_line) > 100:
            continue

        if "@" in line:
            continue

        if re.search(r"\d{7,}", line):
            continue

        if clean_line.lower() in ignored_words:
            continue

        return clean_line

    return None


# ============================================================
# SKILLS
# ============================================================

def extract_skills(text):
    """Identify known technical skills."""

    text_lower = text.lower()

    found_skills = []

    for skill in SKILLS:

        # Escape special regex characters
        pattern = r"(?<!\w)" + re.escape(skill.lower()) + r"(?!\w)"

        if re.search(pattern, text_lower):

            found_skills.append(skill)

    return sorted(set(found_skills))


# ============================================================
# EXPERIENCE
# ============================================================

def extract_experience(text):
    """
    Extract experience information from an Experience/Employment/
    Internship section.
    """

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    experience = []

    inside_experience = False

    experience_headers = [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "employment history",
        "internship",
        "internships",
        "work history",
    ]

    stop_headers = [
        "education",
        "skills",
        "technical skills",
        "projects",
        "academic projects",
        "personal projects",
        "certifications",
        "achievements",
        "awards",
        "interests",
        "references",
    ]

    for line in lines:

        lower = line.lower().strip()

        # Start experience section
        if lower in experience_headers:
            inside_experience = True
            continue

        # Stop when another major section begins
        if inside_experience and lower in stop_headers:
            break

        if inside_experience:

            # Ignore extremely long lines
            if len(line) > 250:
                continue

            experience.append(line)

    # If section-based extraction failed,
    # try detecting explicit years of experience.
    if not experience:

        patterns = [
            r"(\d+(?:\.\d+)?)\+?\s+years?\s+of\s+experience",
            r"(\d+(?:\.\d+)?)\+?\s+years?\s+experience",
        ]

        for pattern in patterns:

            matches = re.findall(
                pattern,
                text,
                re.IGNORECASE
            )

            if matches:

                return [
                    f"{match} years of experience"
                    for match in matches
                ]

    return experience[:20]


# ============================================================
# PROJECTS
# ============================================================

def extract_projects(text):
    """
    Extract project names from the Projects section.

    A project normally begins with a bullet point or a line
    ending with ':'.
    """

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    projects = []

    inside_projects = False

    project_headers = [
        "projects",
        "project",
        "academic projects",
        "personal projects",
        "key projects",
    ]

    stop_headers = [
        "experience",
        "work experience",
        "professional experience",
        "education",
        "skills",
        "technical skills",
        "certifications",
        "achievements",
        "awards",
    ]

    current_project = None

    for line in lines:

        lower = line.lower().strip()

        # Start Projects section
        if lower in project_headers:

            inside_projects = True
            continue

        # Stop Projects section
        if inside_projects and lower in stop_headers:
            break

        if not inside_projects:
            continue

        # Remove common bullet characters
        cleaned = re.sub(
            r"^[•●▪◦\-\*\u200b\s]+",
            "",
            line
        ).strip()

        if not cleaned:
            continue

        # Detect project heading
        #
        # Example:
        # Student Management System(Java,JDBC):
        #
        # Expense Tracker Application:
        #
        if cleaned.endswith(":"):

            project_name = cleaned.rstrip(":").strip()

            if current_project:
                projects.append(current_project)

            current_project = project_name

        # Detect another common project-heading pattern
        elif (
            len(cleaned) <= 100
            and (
                "(" in cleaned
                or "application" in cleaned.lower()
                or "system" in cleaned.lower()
            )
            and not cleaned.lower().startswith(
                (
                    "developed",
                    "implemented",
                    "created",
                    "designed",
                    "used",
                    "built",
                    "worked",
                )
            )
        ):

            if current_project:
                projects.append(current_project)

            current_project = cleaned

        # Description line
        else:

            # If no project heading has been found yet,
            # treat the first meaningful line as a project.
            if current_project is None:

                current_project = cleaned

    # Add final project
    if current_project:
        projects.append(current_project)

    # Clean duplicate projects
    cleaned_projects = []

    for project in projects:

        project = project.strip()

        if project and project not in cleaned_projects:
            cleaned_projects.append(project)

    return cleaned_projects[:10]


# ============================================================
# KEYWORDS
# ============================================================

def extract_keywords(text):
    """
    Extract useful ATS-oriented keywords using spaCy.

    Removes:
    - URLs
    - emails
    - phone numbers
    - stop words
    - very short words
    - noisy generic words
    """

    # Remove URLs
    clean_text = re.sub(
        r"https?://\S+|www\.\S+",
        " ",
        text,
        flags=re.IGNORECASE
    )

    # Remove email addresses
    clean_text = re.sub(
        r"\S+@\S+",
        " ",
        clean_text
    )

    # Remove phone numbers
    clean_text = re.sub(
        r"\+?\d[\d\s().-]{7,}\d",
        " ",
        clean_text
    )

    doc = nlp(clean_text)

    keywords = []

    ignored_words = {
        "resume",
        "curriculum",
        "vitae",
        "profile",
        "summary",
        "detail",
        "details",
        "year",
        "student",
        "candidate",
        "linkedin",
        "email",
        "phone",
        "mobile",
        "number",
        "information",
        "experience",
        "project",
        "projects",
        "work",
        "education",
        "skill",
        "skills",
        "application",
    }

    for token in doc:

        if token.is_stop:
            continue

        if token.is_punct:
            continue

        if token.is_space:
            continue

        if token.like_url:
            continue

        if token.like_email:
            continue

        if token.like_num:
            continue

        if token.pos_ not in [
            "NOUN",
            "PROPN",
            "ADJ",
        ]:
            continue

        word = token.lemma_.lower().strip()

        if len(word) < 3:
            continue

        if word in ignored_words:
            continue

        if not re.match(
            r"^[a-zA-Z][a-zA-Z+#.-]*$",
            word
        ):
            continue

        keywords.append(word)

    # Add technical skills explicitly.
    # This makes the keyword list useful for future ATS matching.
    skill_keywords = extract_skills(text)

    keywords.extend(skill_keywords)

    # Remove duplicates while preserving order
    unique_keywords = list(
        dict.fromkeys(keywords)
    )

    return unique_keywords[:40]


# ============================================================
# COMPLETE RESUME PARSER
# ============================================================

def parse_resume(file_path):
    """Run the complete resume parsing pipeline."""

    text = extract_resume_text(file_path)

    parsed_data = {
        "extracted_text": text,

        "name": extract_name(text),

        "email": extract_email(text),

        "phone": extract_phone(text),

        "skills": extract_skills(text),

        "experience": extract_experience(text),

        "projects": extract_projects(text),

        "keywords": extract_keywords(text),
    }

    return parsed_data